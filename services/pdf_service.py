# services/pdf_service.py
import os
import glob
import uuid
import platform
import logging
import tempfile
import subprocess
import threading
from io import BytesIO
from pathlib import Path
from docx import Document

from helpers.docx_helpers import (
    replace_text_in_paragraph,
    replace_text_in_tables
)

logger = logging.getLogger(__name__)

# ==================== PRINT JOB TRACKER ====================
print_jobs = {}
print_jobs_lock = threading.Lock()


def create_print_job(job_id):
    """Register a new print job"""
    with print_jobs_lock:
        print_jobs[job_id] = {'cancelled': False, 'process': None}


def set_print_process(job_id, proc):
    """Attach the LibreOffice subprocess to the job so it can be killed"""
    with print_jobs_lock:
        if job_id in print_jobs:
            print_jobs[job_id]['process'] = proc


def is_job_cancelled(job_id):
    """Thread-safe check: has this job been cancelled?"""
    with print_jobs_lock:
        return print_jobs.get(job_id, {}).get('cancelled', False)


def cleanup_print_job(job_id):
    """Remove a job from the tracker"""
    with print_jobs_lock:
        print_jobs.pop(job_id, None)


def cancel_print_job(job_id):
    """
    Mark a job as cancelled and kill its LibreOffice process if running.
    Returns (success, message)
    """
    with print_jobs_lock:
        if job_id not in print_jobs:
            return False, 'Job not found or already completed'

        print_jobs[job_id]['cancelled'] = True
        proc = print_jobs[job_id].get('process')

        if proc and proc.poll() is None:
            try:
                proc.kill()
                logger.info(f'[CancelPrint] Killed LibreOffice process for job {job_id}')
            except Exception as kill_err:
                logger.warning(f'[CancelPrint] Could not kill process: {kill_err}')

    return True, 'Print job cancelled successfully'


# ==================== RAM DISK ====================

def get_ram_disk_path():
    """Get RAM disk path for fastest I/O."""
    if platform.system() == 'Linux' and os.path.exists('/dev/shm'):
        ram_path = '/dev/shm/lo_pdf_work'
        os.makedirs(ram_path, exist_ok=True)
        return ram_path
    return None


# ==================== FIND LIBREOFFICE ====================

def find_libreoffice_executable():
    """Find the LibreOffice executable on Windows, Linux, or macOS."""
    system = platform.system()

    if system == 'Windows':
        search_bases = [
            os.environ.get('PROGRAMFILES', r'C:\Program Files'),
            os.environ.get('PROGRAMFILES(X86)', r'C:\Program Files (x86)'),
            r'C:\Program Files',
            r'C:\Program Files (x86)',
        ]
        for base in search_bases:
            if not base:
                continue
            pattern = os.path.join(
                base, 'LibreOffice*', 'program', 'soffice.exe'
            )
            matches = glob.glob(pattern)
            if matches:
                return matches[0]
        return 'soffice.exe'

    elif system == 'Darwin':
        mac_path = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
        if os.path.exists(mac_path):
            return mac_path
        return 'soffice'

    else:
        return 'libreoffice'


# ==================== BATCH CONVERT DOCX → PDF ====================

def convert_all_docx_to_pdfs_batch(docx_files_dict, replacements, job_id=None):
    """
    Fill all DOCX templates and convert to PDF in a single LibreOffice call.
    Supports cancellation via job_id.
    Returns dict of {file_stem: pdf_bytes}
    """
    import time
    libreoffice_cmd = find_libreoffice_executable()

    ram_path = get_ram_disk_path()
    session_id = uuid.uuid4().hex[:8]

    if ram_path:
        work_dir = os.path.join(ram_path, session_id)
    else:
        work_dir = os.path.join(tempfile.gettempdir(), f'lo_work_{session_id}')

    os.makedirs(work_dir, exist_ok=True)
    profile_dir = os.path.join(work_dir, 'lo_profile')
    os.makedirs(profile_dir, exist_ok=True)

    logger.info(f'[BatchPDF] Work dir: {work_dir}')

    try:
        # ── Check cancellation before starting ───────────────────────
        if job_id and is_job_cancelled(job_id):
            raise RuntimeError('CANCELLED')

        # ── Write font substitution profile ──────────────────────────
        registry_dir = os.path.join(profile_dir, 'user')
        os.makedirs(registry_dir, exist_ok=True)

        font_config = '''<?xml version="1.0" encoding="UTF-8"?>
<oor:items xmlns:oor="http://openoffice.org/2001/registry"
           xmlns:xs="http://www.w3.org/2001/XMLSchema"
           xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
    <item oor:path="/org.openoffice.VCL/SubstFonts">
        <prop oor:name="Calibri" oor:op="fuse">
            <value>Liberation Sans</value>
        </prop>
        <prop oor:name="Calibri (Body)" oor:op="fuse">
            <value>Liberation Sans</value>
        </prop>
        <prop oor:name="Calibri Light" oor:op="fuse">
            <value>Liberation Sans</value>
        </prop>
        <prop oor:name="Times New Roman" oor:op="fuse">
            <value>Liberation Serif</value>
        </prop>
        <prop oor:name="Arial" oor:op="fuse">
            <value>Liberation Sans</value>
        </prop>
        <prop oor:name="Arial Narrow" oor:op="fuse">
            <value>Liberation Sans Narrow</value>
        </prop>
        <prop oor:name="Cambria" oor:op="fuse">
            <value>Liberation Serif</value>
        </prop>
        <prop oor:name="Cambria (Body)" oor:op="fuse">
            <value>Liberation Serif</value>
        </prop>
        <prop oor:name="Courier New" oor:op="fuse">
            <value>Liberation Mono</value>
        </prop>
        <prop oor:name="Georgia" oor:op="fuse">
            <value>Liberation Serif</value>
        </prop>
        <prop oor:name="Verdana" oor:op="fuse">
            <value>Liberation Sans</value>
        </prop>
        <prop oor:name="Tahoma" oor:op="fuse">
            <value>Liberation Sans</value>
        </prop>
        <prop oor:name="Mangal" oor:op="fuse">
            <value>Liberation Sans</value>
        </prop>
    </item>
</oor:items>'''

        registry_file = os.path.join(registry_dir, 'registrymodifications.xcu')
        with open(registry_file, 'w', encoding='utf-8') as f:
            f.write(font_config)

        # ── Check cancellation before filling DOCX ───────────────────
        if job_id and is_job_cancelled(job_id):
            raise RuntimeError('CANCELLED')

        # ── Fill ALL DOCX templates ───────────────────────────────────
        docx_written = []

        for file_stem, docx_path in docx_files_dict.items():
            if job_id and is_job_cancelled(job_id):
                raise RuntimeError('CANCELLED')

            try:
                logger.info(f'[BatchPDF] Filling: {file_stem}')

                doc = Document(str(docx_path))

                for paragraph in doc.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

                if doc.tables:
                    replace_text_in_tables(doc.tables, replacements)

                for section in doc.sections:
                    for para in section.header.paragraphs:
                        replace_text_in_paragraph(para, replacements)
                    if section.header.tables:
                        replace_text_in_tables(section.header.tables, replacements)
                    for para in section.footer.paragraphs:
                        replace_text_in_paragraph(para, replacements)
                    if section.footer.tables:
                        replace_text_in_tables(section.footer.tables, replacements)

                out_docx = os.path.join(work_dir, f'{file_stem}.docx')
                doc.save(out_docx)
                docx_written.append((file_stem, out_docx))
                logger.info(f'[BatchPDF] Saved: {file_stem}.docx')

            except RuntimeError:
                raise  # Re-raise CANCELLED
            except Exception as e:
                logger.error(f'[BatchPDF] Error filling {file_stem}: {e}')
                continue

        if not docx_written:
            raise RuntimeError('No DOCX files could be filled/written')

        # ── Check cancellation before LibreOffice call ────────────────
        if job_id and is_job_cancelled(job_id):
            raise RuntimeError('CANCELLED')

        # ── Build profile URL ─────────────────────────────────────────
        if platform.system() == 'Windows':
            profile_path_forward = profile_dir.replace('\\', '/')
            if not profile_path_forward.startswith('/'):
                profile_url = f'file:///{profile_path_forward}'
            else:
                profile_url = f'file://{profile_path_forward}'
        else:
            profile_url = f'file://{profile_dir}'

        logger.info(f'[BatchPDF] Profile URL: {profile_url}')

        # ── Call LibreOffice ONCE for ALL files ───────────────────────
        all_docx_paths = [path for _, path in docx_written]

        cmd = [
            libreoffice_cmd,
            f'-env:UserInstallation={profile_url}',
            '--headless',
            '--norestore',
            '--nofirststartwizard',
            '--convert-to', 'pdf',
            '--outdir', work_dir,
        ] + all_docx_paths

        logger.info(f'[BatchPDF] Running LO for {len(all_docx_paths)} files')

        env = dict(os.environ)
        if platform.system() != 'Windows':
            env['HOME'] = work_dir

        try:
            start_time = time.time()

            # ── Launch as Popen so we can kill it if cancelled ────────
            proc = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                env=env
            )

            # Register process so cancel endpoint can kill it
            if job_id:
                set_print_process(job_id, proc)

            # Poll every 0.5s so we can react to cancellation quickly
            stdout = ''
            stderr = ''
            while True:
                try:
                    stdout, stderr = proc.communicate(timeout=0.5)
                    break  # Process finished normally
                except subprocess.TimeoutExpired:
                    if job_id and is_job_cancelled(job_id):
                        proc.kill()
                        proc.communicate()  # Clean up
                        raise RuntimeError('CANCELLED')

            elapsed = time.time() - start_time
            logger.info(f'[BatchPDF] LO took {elapsed:.2f}s')

            # Filter harmless warnings from stderr
            if stderr:
                harmless_warnings = [
                    'Could not find platform independent libraries',
                    'Could not find platform dependent libraries',
                    'Consider setting $PYTHONHOME',
                    'ImportError: No module named site',
                    'cannot set LC_',
                    'javaldx: Could not find a Java',
                ]
                stderr_lines = stderr.strip().splitlines()
                real_errors = [
                    line for line in stderr_lines
                    if line.strip()
                    and not any(w in line for w in harmless_warnings)
                ]
                if real_errors:
                    logger.warning(
                        f'[BatchPDF] stderr: {chr(10).join(real_errors)}'
                    )

            if stdout:
                logger.info(f'[BatchPDF] stdout: {stdout}')

            result_returncode = proc.returncode
            result_stdout = stdout
            result_stderr = stderr

        except RuntimeError:
            raise  # Re-raise CANCELLED
        except FileNotFoundError:
            raise RuntimeError(
                f"LibreOffice not found at '{libreoffice_cmd}'. "
                f"Install from https://www.libreoffice.org/download/"
            )

        if result_returncode != 0:
            raise RuntimeError(
                f'LibreOffice error {result_returncode}.\n'
                f'stderr: {result_stderr}\nstdout: {result_stdout}'
            )

        # ── Check cancellation before reading PDFs ────────────────────
        if job_id and is_job_cancelled(job_id):
            raise RuntimeError('CANCELLED')

        # ── Read ALL generated PDFs into memory ───────────────────────
        pdf_results = {}
        missing = []

        for file_stem, _ in docx_written:
            pdf_path = os.path.join(work_dir, f'{file_stem}.pdf')

            if os.path.exists(pdf_path):
                with open(pdf_path, 'rb') as f:
                    pdf_results[file_stem] = f.read()
                logger.info(
                    f'[BatchPDF] ✅ Read: {file_stem}.pdf '
                    f'({len(pdf_results[file_stem])} bytes)'
                )
            else:
                missing.append(file_stem)
                logger.warning(f'[BatchPDF] ❌ Missing: {file_stem}.pdf')

        if not pdf_results:
            raise RuntimeError(
                f'LibreOffice produced no PDF output.\n'
                f'Missing: {missing}\n'
                f'stdout: {result_stdout}\nstderr: {result_stderr}'
            )

        if missing:
            logger.warning(f'[BatchPDF] Some PDFs missing: {missing}')

        logger.info(
            f'[BatchPDF] ✅ Done: {len(pdf_results)}/{len(docx_written)} converted'
        )
        return pdf_results

    finally:
        try:
            import shutil
            shutil.rmtree(work_dir, ignore_errors=True)
            logger.info(f'[BatchPDF] Cleaned up: {work_dir}')
        except Exception as cleanup_err:
            logger.warning(f'[BatchPDF] Cleanup warning: {cleanup_err}')


# ==================== MERGE PDFs ====================

def merge_pdfs_bytes(pdf_bytes_list):
    """Merge a list of PDF byte-strings into one. Everything in memory."""
    from pypdf import PdfWriter, PdfReader

    writer = PdfWriter()

    for pdf_bytes in pdf_bytes_list:
        reader = PdfReader(BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)

    output = BytesIO()
    writer.write(output)
    output.seek(0)

    logger.info(f'[MergePDF] Merged {len(pdf_bytes_list)} PDFs successfully')
    return output.read()