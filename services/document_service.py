# services/document_service.py
import re
import logging
import zipfile
from io import BytesIO
from pathlib import Path
from datetime import datetime, timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from helpers.docx_helpers import (
    replace_text_in_paragraph,
    replace_text_in_tables,
    get_templates
)
from helpers.text_helpers import (
    resolve_he_she,
    create_safe_folder_name
)
from helpers.html_helpers import (
    process_paragraph_html,
    process_table_html
)
from services.template_service import resolve_template_folder

logger = logging.getLogger(__name__)


def prepare_replacements(replacements):
    """
    Ensure all computed fields like HE_SHE are correctly resolved.
    Call this before any document rendering/preview/print.
    """
    replacements = dict(replacements)

    if 'WIFE_OF' not in replacements:
        replacements['WIFE_OF'] = ''
    if 'SPOUSE_NAME1' not in replacements:
        replacements['SPOUSE_NAME1'] = ''

    # Normalize SON-DAUGHTER to lowercase
    if 'SON-DAUGHTER' in replacements and replacements['SON-DAUGHTER']:
        replacements['SON-DAUGHTER'] = str(replacements['SON-DAUGHTER']).lower().strip()

    # Always recalculate HE_SHE from source fields
    son_daughter = replacements.get('SON-DAUGHTER', '').strip()
    gender = replacements.get('GENDER_UPDATE', '').strip()
    existing = replacements.get('HE_SHE', '').strip()

    replacements['HE_SHE'] = resolve_he_she(
        son_daughter=son_daughter,
        gender=gender,
        existing=existing
    )

    logger.info(
        f'[PrepareReplacements] HE_SHE={repr(replacements["HE_SHE"])} | '
        f'SON-DAUGHTER={repr(son_daughter)} | '
        f'GENDER_UPDATE={repr(gender)}'
    )

    return replacements


def generate_documents_to_memory(draft, template_config):
    """Generate documents in memory and return as ZIP buffer"""
    try:
        template_type = draft.template_type
        replacements = draft.replacements or {}
        preview_data = draft.preview_data or {}

        if template_type not in template_config:
            return {'success': False, 'message': 'Invalid template type'}

        # Prepare replacements
        replacements = prepare_replacements(replacements)

        # Resolve template folder
        template_folder_str = resolve_template_folder(
            template_type, template_config, preview_data, replacements
        )

        if not template_folder_str:
            return {'success': False, 'message': 'Template folder not configured'}

        template_folder = Path(template_folder_str)
        templates = get_templates(template_folder)

        if not templates:
            return {
                'success': False,
                'message': f'No templates found in {template_folder}'
            }

        if template_type == 'minor_template':
            folder_name_source = replacements.get('FATHER-MOTHER_NAME', '')
            if not folder_name_source:
                folder_name_source = replacements.get('OLD_NAME', 'unnamed')
        else:
            folder_name_source = replacements.get('OLD_NAME', 'unnamed')

        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for template_name in templates:
                input_path = template_folder / template_name

                doc = Document(str(input_path))

                for paragraph in doc.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

                if doc.tables:
                    replace_text_in_tables(doc.tables, replacements)

                for section in doc.sections:
                    header = section.header
                    for paragraph in header.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                    if header.tables:
                        replace_text_in_tables(header.tables, replacements)

                for section in doc.sections:
                    footer = section.footer
                    for paragraph in footer.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                    if footer.tables:
                        replace_text_in_tables(footer.tables, replacements)

                doc_buffer = BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                original_name = Path(template_name).stem
                safe_folder = create_safe_folder_name(folder_name_source)
                output_filename = f"{original_name} {safe_folder}.docx"

                zipf.writestr(output_filename, doc_buffer.read())

        return {
            'success': True,
            'zip_buffer': zip_buffer,
            'file_count': len(templates)
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        return {'success': False, 'message': str(e)}


def get_cd_document_content(template_folder, replacements):
    """Get the content of CD.docx document with updated values"""
    try:
        template_folder_path = Path(template_folder)

        cd_file = None
        if template_folder_path.exists():
            for file in template_folder_path.iterdir():
                if (file.is_file()
                        and file.suffix.lower() == '.docx'
                        and file.stem.upper() == 'CD'
                        and not file.name.startswith('~$')):
                    cd_file = file
                    break

        if not cd_file:
            return f"""
                <div class="text-center py-4">
                    <i class="bi bi-file-earmark-x text-muted" style="font-size: 3rem;"></i>
                    <p class="text-muted mt-2">CD.docx not found in template folder.</p>
                    <small class="text-muted">Looking in: {template_folder}</small>
                </div>
            """

        doc = Document(str(cd_file))

        clean_replacements = {
            key: str(value).strip() if value is not None else ''
            for key, value in replacements.items()
        }

        html_parts = ['<div class="cd-content-wrapper">']

        for para in doc.paragraphs:
            text = para.text

            if not text.strip():
                html_parts.append('<p class="cd-paragraph">&nbsp;</p>')
                continue

            for key, value in clean_replacements.items():
                if key in text:
                    if value:
                        highlighted = (
                            f'<span class="cd-replaced-value">{value}</span>'
                        )
                        text = text.replace(key, highlighted)
                    else:
                        text = text.replace(key, '')

            text = re.sub(r'\s+', ' ', text)
            text = re.sub(r'\s+,', ',', text)
            text = re.sub(r'\s+\.', '.', text)
            text = re.sub(r'\s+;', ';', text)
            text = re.sub(r'\(\s+', '(', text)
            text = re.sub(r'\s+\)', ')', text)
            text = text.strip()

            if not text:
                continue

            style_class = 'cd-paragraph'
            if para.style and para.style.name:
                if 'Heading' in para.style.name:
                    style_class = 'cd-heading'
                elif 'Title' in para.style.name:
                    style_class = 'cd-title'

            alignment_style = ''
            if para.alignment:
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    alignment_style = ' style="text-align: center;"'
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    alignment_style = ' style="text-align: right;"'
                elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    alignment_style = ' style="text-align: justify;"'

            html_parts.append(
                f'<p class="{style_class}"{alignment_style}>{text}</p>'
            )

        for table in doc.tables:
            html_parts.append('<table class="table table-bordered cd-table">')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    cell_text = cell.text
                    for key, value in clean_replacements.items():
                        if key in cell_text:
                            if value:
                                highlighted = (
                                    f'<span class="cd-replaced-value">{value}</span>'
                                )
                                cell_text = cell_text.replace(key, highlighted)
                            else:
                                cell_text = cell_text.replace(key, '')

                    cell_text = re.sub(r'\s+', ' ', cell_text).strip()
                    cell_text = re.sub(r'\s+,', ',', cell_text)
                    cell_text = re.sub(r'\s+\.', '.', cell_text)
                    html_parts.append(f'<td>{cell_text}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')

        html_parts.append('</div>')

        css = '''
        <style>
            .cd-content-wrapper {
                font-family: 'Times New Roman', serif;
                font-size: 14px;
                line-height: 1.8;
                color: #333;
                background: white;
                padding: 1.5rem;
                border-radius: 8px;
            }
            .cd-paragraph {
                margin-bottom: 0.75rem;
                text-align: justify;
                line-height: 1.8;
            }
            .cd-heading {
                font-weight: bold;
                font-size: 16px;
                margin-bottom: 1rem;
                text-align: center;
                color: #1a202c;
            }
            .cd-title {
                font-weight: bold;
                font-size: 18px;
                margin-bottom: 1rem;
                text-align: center;
                text-transform: uppercase;
                color: #1a202c;
            }
            .cd-replaced-value {
                background-color: #d4edda;
                color: #155724;
                padding: 2px 6px;
                border-radius: 4px;
                font-weight: 600;
                border: 1px solid #c3e6cb;
            }
            .cd-table {
                font-size: 13px;
                margin: 1rem 0;
                width: 100%;
            }
            .cd-table td {
                padding: 0.75rem;
                vertical-align: top;
                border: 1px solid #dee2e6;
            }
            .cd-table tr:hover {
                background-color: #f8f9fa;
            }
        </style>
        '''

        return css + ''.join(html_parts)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"""
            <div class="text-center py-4">
                <i class="bi bi-exclamation-triangle text-danger" 
                   style="font-size: 3rem;"></i>
                <p class="text-danger mt-2">Error loading CD document</p>
                <small class="text-muted">{str(e)}</small>
            </div>
        """