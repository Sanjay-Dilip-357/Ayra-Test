# routes/documents.py
import re
import logging
import zipfile
from io import BytesIO
from pathlib import Path
from datetime import datetime, timezone
from flask import (
    Blueprint, request, jsonify,
    session, send_file, render_template
)
from models import db, Draft
from config import TEMPLATE_CONFIG, RELATION_MAPPING, CAST_OPTIONS
from routes.auth import login_required
from services.document_service import (
    generate_documents_to_memory,
    prepare_replacements
)
from services.phone_service import (
    mark_phones_as_used,
    get_phone_stats
)
from services.template_service import get_all_template_info
from helpers.text_helpers import (
    to_uppercase,
    to_uppercase_preserve_alias,
    format_date_to_ddmmyyyy,
    generate_email_from_name,
    create_safe_folder_name,
    resolve_he_she
)
from helpers.docx_helpers import (
    get_templates,
    replace_text_in_paragraph,
    replace_text_in_tables
)
from docx import Document

logger = logging.getLogger(__name__)

documents_bp = Blueprint('documents', __name__)


@documents_bp.route('/namechangeservice')
@login_required
def namechangeservice():
    """Name change service page"""
    template_info = get_all_template_info(TEMPLATE_CONFIG)
    phone_stats = get_phone_stats()

    return render_template(
        'namechangeservice.html',
        template_info=template_info,
        template_config=TEMPLATE_CONFIG,
        cast_options=CAST_OPTIONS,
        phone_stats=phone_stats,
        user_name=session.get('user_name')
    )


@documents_bp.route('/get_template_config/<template_type>')
@login_required
def get_template_config_route(template_type):
    """Get template config for a type"""
    if template_type in TEMPLATE_CONFIG:
        config = TEMPLATE_CONFIG[template_type]
        templates = get_templates(config['folder'])
        return jsonify({
            'success': True,
            'config': config,
            'templates': templates,
            'count': len(templates),
            'folder': config['folder']
        })
    return jsonify({'success': False, 'message': 'Template type not found'})


@documents_bp.route('/get_templates_by_relation/<template_type>/<relation>')
@login_required
def get_templates_by_relation(template_type, relation):
    """Get templates filtered by relation"""
    if template_type not in TEMPLATE_CONFIG:
        return jsonify({'success': False, 'message': 'Template type not found'})

    config = TEMPLATE_CONFIG[template_type]

    if template_type in ['major_template', 'religion_template']:
        if relation.lower() == 'd':
            if 'unmarried_subfolder' in config:
                folder = config['unmarried_subfolder']
                templates = get_templates(folder)
                return jsonify({
                    'success': True,
                    'templates': templates,
                    'count': len(templates),
                    'folder': folder,
                    'folder_type': 'unmarried'
                })

    folder = config['folder']
    templates = get_templates(folder)
    return jsonify({
        'success': True,
        'templates': templates,
        'count': len(templates),
        'folder': folder,
        'folder_type': 'main'
    })


@documents_bp.route('/generate_email', methods=['POST'])
@login_required
def generate_email_route():
    """Generate email from name"""
    name = request.json.get('name', '')
    email = generate_email_from_name(name)
    return jsonify({'success': True, 'email': email})


@documents_bp.route('/preview', methods=['POST'])
@login_required
def preview_document():
    """Generate preview"""
    try:
        template_type = request.form.get('template_type', '').strip()

        if template_type not in TEMPLATE_CONFIG:
            return jsonify({'success': False, 'message': 'Invalid template type!'})

        template_config = TEMPLATE_CONFIG[template_type]

        relation_input = request.form.get('relation', '').strip().lower()
        update_relation = RELATION_MAPPING.get(relation_input, '')

        if template_type in ['major_template', 'religion_template'] and relation_input == 'd':
            spouse_name_check = (
                request.form.get('spouse_name', '').strip()
                if relation_input == 'd/w' else ''
            )

            if not spouse_name_check and 'unmarried_subfolder' in template_config:
                template_folder = template_config['unmarried_subfolder']
                folder_type = 'unmarried'
            else:
                template_folder = template_config['folder']
                folder_type = 'main'
        else:
            template_folder = template_config['folder']
            folder_type = 'main'

        templates = get_templates(template_folder)

        if not templates:
            return jsonify({'success': False, 'message': 'No template files found!'})

        old_name_raw = request.form.get('old_name', '').strip()
        old_name = to_uppercase_preserve_alias(old_name_raw)

        new_name = to_uppercase(request.form.get('new_name', ''))
        gender = request.form.get('gender_update', '').strip()

        if not old_name:
            return jsonify({'success': False, 'message': 'Name field is required!'})

        base_old_name = re.split(
            r'\s+alias\s+', old_name, flags=re.IGNORECASE
        )[0].strip()

        phone_update = request.form.get('phone_update', '').strip()
        witness_phone1 = request.form.get('witness_phone1', '').strip()
        witness_phone2 = request.form.get('witness_phone2', '').strip()

        phones_used = [p for p in [phone_update, witness_phone1, witness_phone2] if p]

        replacements = {
            'OLD_NAME': old_name,
            'NEW_NAME': new_name if new_name else base_old_name,
            'UPDATE_ADDRESS': to_uppercase(request.form.get('update_address', '')),
            'GENDER_UPDATE': to_uppercase(gender),
            'CAST_UPDATE': to_uppercase(request.form.get('cast_update', '')),
            'PHONE_UPDATE': phone_update,
            'EMAIL_UPDATE': request.form.get('email_update', '').strip(),
            'NUM_DATE': request.form.get('num_date', '').strip(),
            'ALPHA_DATE': request.form.get('alpha_date', '').strip(),
            'WITNESS_NAME1': to_uppercase(request.form.get('witness_name1', '')),
            'WITNESS_ADDRESS1': to_uppercase(request.form.get('witness_address1', '')),
            'WITNESS_PHONE1': witness_phone1,
            'WITNESS_NAME2': to_uppercase(request.form.get('witness_name2', '')),
            'WITNESS_ADDRESS2': to_uppercase(request.form.get('witness_address2', '')),
            'WITNESS_PHONE2': witness_phone2,
        }

        if relation_input == 'd/w':
            father_name = to_uppercase(request.form.get('father_name', ''))
            spouse_name = to_uppercase(request.form.get('spouse_name', ''))
            replacements.update({
                'UPDATE_RELATION': 'D/o',
                'FATHER-SPOUSE_NAME': father_name,
                'WIFE_OF': ' W/o ',
                'SPOUSE_NAME1': spouse_name,
            })
        else:
            fatherspouse_name = to_uppercase(
                request.form.get('fatherspouse_name', '')
            )
            replacements.update({
                'UPDATE_RELATION': update_relation,
                'FATHER-SPOUSE_NAME': fatherspouse_name,
                'WIFE_OF': '',
                'SPOUSE_NAME1': '',
            })

        son_daughter = None
        folder_name_source = base_old_name

        if template_type == 'minor_template':
            child_dob_raw = request.form.get('child_dob', '').strip()
            child_dob_formatted = format_date_to_ddmmyyyy(child_dob_raw)

            son_daughter = request.form.get('son_daughter', '').strip()
            fathermother_name = to_uppercase(
                request.form.get('fathermother_name', '')
            )

            if relation_input == 'd/w':
                guardian_father_name = to_uppercase(
                    request.form.get('guardian_father_name', '')
                )
                guardian_spouse_name = to_uppercase(
                    request.form.get('guardian_spouse_name', '')
                )
                replacements.update({
                    'UPDATE_RELATION': 'D/o',
                    'FATHER-SPOUSE_NAME': guardian_father_name,
                    'WIFE_OF': ' W/o ',
                    'SPOUSE_NAME1': guardian_spouse_name,
                })

            replacements.update({
                'UPDATE_AGE': request.form.get('update_age', '').strip(),
                'FATHER-MOTHER_NAME': fathermother_name,
                'SON-DAUGHTER': son_daughter,
                'CHILD_DOB': child_dob_formatted,
                'BIRTH_PLACE': to_uppercase(request.form.get('birth_place', '')),
            })

            if fathermother_name:
                folder_name_source = fathermother_name
            else:
                folder_name_source = base_old_name

        replacements['HE_SHE'] = resolve_he_she(
            son_daughter=son_daughter,
            gender=gender,
            existing=None
        )

        session['preview_data'] = {
            'template_type': template_type,
            'template_folder': str(template_folder),
            'templates': templates,
            'replacements': replacements,
            'folder_name_source': folder_name_source,
            'relation_input': relation_input,
            'folder_type': folder_type,
            'phones_used': phones_used
        }

        display_replacements = {
            k: v for k, v in replacements.items()
            if v and str(v).strip()
        }

        return jsonify({
            'success': True,
            'message': 'Preview generated successfully!',
            'template_type': template_type,
            'template_name': template_config['name'],
            'template_count': len(templates),
            'replacements': display_replacements,
            'folder_named_by': (
                'Father/Mother Name'
                if template_type == 'minor_template'
                else 'Applicant Name'
            ),
            'used_unmarried_folder': folder_type == 'unmarried'
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})


@documents_bp.route('/update_preview', methods=['POST'])
@login_required
def update_preview():
    """Update preview data"""
    try:
        if 'preview_data' not in session:
            return jsonify({'success': False, 'message': 'No preview data found.'})

        updated_replacements = request.json.get('replacements', {})

        for key, value in updated_replacements.items():
            if key in session['preview_data']['replacements']:
                session['preview_data']['replacements'][key] = value

        session.modified = True

        return jsonify({'success': True, 'message': 'Preview updated!'})

    except Exception as e:
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})


@documents_bp.route('/generate', methods=['POST'])
@login_required
def generate_document():
    """Generate and download documents directly"""
    try:
        if 'preview_data' not in session:
            return jsonify({'success': False, 'message': 'No preview data found.'})

        preview_data = session['preview_data']

        template_type = preview_data['template_type']
        template_folder = Path(preview_data['template_folder'])
        templates = preview_data['templates']
        replacements = preview_data['replacements']

        phones_used = preview_data.get('phones_used', [])
        mark_phones_as_used(phones_used)

        if template_type == 'minor_template':
            folder_name_source = replacements.get('FATHER-MOTHER_NAME', '')
            if not folder_name_source:
                folder_name_source = preview_data.get(
                    'folder_name_source',
                    replacements.get('OLD_NAME', 'unnamed')
                )
        else:
            folder_name_source = preview_data.get(
                'folder_name_source',
                replacements.get('OLD_NAME', 'unnamed')
            )

        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for template_name in templates:
                input_path = template_folder / template_name

                doc = Document(str(input_path))

                for paragraph in doc.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

                if doc.tables:
                    replace_text_in_tables(doc.tables, replacements)

                for section in doc.sections:
                    header = section.header
                    for paragraph in header.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                    if header.tables:
                        replace_text_in_tables(header.tables, replacements)

                for section in doc.sections:
                    footer = section.footer
                    for paragraph in footer.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                    if footer.tables:
                        replace_text_in_tables(footer.tables, replacements)

                doc_buffer = BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                original_name = Path(template_name).stem
                safe_folder = create_safe_folder_name(folder_name_source)
                output_filename = f'{original_name} {safe_folder}.docx'

                zipf.writestr(output_filename, doc_buffer.read())

        session.pop('preview_data', None)

        zip_buffer.seek(0)
        filename = (
            f"{create_safe_folder_name(folder_name_source)}"
        )

        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})


@documents_bp.route('/save_draft', methods=['POST'])
@login_required
def save_draft():
    """Save draft from preview session"""
    try:
        if 'preview_data' not in session:
            return jsonify({'success': False, 'message': 'No preview data found.'})

        preview_data = session['preview_data']
        user_id = session.get('user_id')

        template_config = TEMPLATE_CONFIG.get(preview_data['template_type'], {})

        phones_used = preview_data.get('phones_used', [])
        if phones_used:
            mark_phones_as_used(phones_used)
            logger.info(f'Marked phones as used from save_draft: {phones_used}')

        new_draft = Draft(
            user_id=user_id,
            template_type=preview_data['template_type'],
            template_name=template_config.get('name', preview_data['template_type']),
            old_name=preview_data['replacements'].get('OLD_NAME', ''),
            replacements=preview_data['replacements'],
            preview_data=preview_data,
            status='draft'
        )

        db.session.add(new_draft)
        db.session.commit()

        session.pop('preview_data', None)

        return jsonify({
            'success': True,
            'message': 'Draft saved successfully!',
            'draft_id': new_draft.id
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})


@documents_bp.route('/api/generate/approved', methods=['GET'])
@login_required
def api_get_approved():
    """Get approved drafts ready for generation"""
    try:
        user_id = session.get('user_id')

        approved = Draft.query.filter_by(
            user_id=user_id, status='approved'
        ).order_by(Draft.approved_at.desc()).all()

        generated = Draft.query.filter_by(
            user_id=user_id, status='generated'
        ).order_by(Draft.generated_at.desc()).all()

        return jsonify({
            'success': True,
            'approved': [d.to_dict() for d in approved],
            'generated': [d.to_dict() for d in generated]
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})


@documents_bp.route('/api/generate/batch', methods=['POST'])
@login_required
def api_generate_batch():
    """Generate documents for multiple approved drafts"""
    try:
        user_id = session.get('user_id')
        draft_ids = request.json.get('draft_ids', [])

        if not draft_ids:
            return jsonify({'success': False, 'message': 'No drafts selected'})

        generated_files = []
        errors = []

        bulk_zip_buffer = BytesIO()

        with zipfile.ZipFile(bulk_zip_buffer, 'w', zipfile.ZIP_DEFLATED) as bulk_zip:
            for draft_id in draft_ids:
                draft = Draft.query.filter_by(
                    id=draft_id, user_id=user_id, status='approved'
                ).first()

                if not draft:
                    errors.append({
                        'draft_id': draft_id,
                        'error': 'Draft not found or not approved'
                    })
                    continue

                try:
                    result = generate_documents_to_memory(draft, TEMPLATE_CONFIG)

                    if result['success']:
                        draft.status = 'generated'
                        draft.generated_at = datetime.now(timezone.utc)
                        draft.generated_files = [
                            f"Generated {result['file_count']} files"
                        ]

                        zip_buffer = result['zip_buffer']
                        zip_buffer.seek(0)

                        folder_name = (
                            draft.old_name.replace(' ', '_')
                            if draft.old_name else f'doc_{draft_id}'
                        )

                        with zipfile.ZipFile(zip_buffer, 'r') as individual_zip:
                            for file_info in individual_zip.filelist:
                                file_data = individual_zip.read(file_info.filename)
                                bulk_zip.writestr(
                                    f'{folder_name}/{file_info.filename}',
                                    file_data
                                )

                        generated_files.append({
                            'draft_id': draft_id,
                            'old_name': draft.old_name,
                            'file_count': result['file_count']
                        })
                    else:
                        errors.append({
                            'draft_id': draft_id,
                            'error': result.get('message')
                        })
                except Exception as e:
                    errors.append({'draft_id': draft_id, 'error': str(e)})

        db.session.commit()

        if generated_files:
            bulk_zip_buffer.seek(0)
            return send_file(
                bulk_zip_buffer,
                mimetype='application/zip',
                as_attachment=True,
                download_name=(
                    f"documents.zip"
                )
            )
        else:
            return jsonify({
                'success': False,
                'message': 'No documents were generated',
                'errors': errors
            })

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@documents_bp.route('/api/generate/single/<draft_id>', methods=['POST'])
@login_required
def api_generate_single(draft_id):
    """Generate a single draft and download"""
    try:
        user_id = session.get('user_id')
        draft = Draft.query.filter_by(
            id=draft_id, user_id=user_id, status='approved'
        ).first()

        if not draft:
            return jsonify({
                'success': False,
                'message': 'Draft not found or not approved'
            })

        result = generate_documents_to_memory(draft, TEMPLATE_CONFIG)

        if not result['success']:
            return jsonify({
                'success': False,
                'message': result.get('message', 'Generation failed')
            })

        draft.status = 'generated'
        draft.generated_at = datetime.now(timezone.utc)
        draft.generated_files = [f"Generated {result['file_count']} files"]
        db.session.commit()

        zip_buffer = result['zip_buffer']
        zip_buffer.seek(0)

        filename = (
            f"{draft.old_name.replace(' ', '_')}_"
        )

        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})