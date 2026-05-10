# routes/admin.py
import re
import html
import uuid
import logging
import zipfile
from io import BytesIO
from pathlib import Path
from datetime import datetime, timezone
from flask import (
    Blueprint, request, jsonify,
    session, render_template, send_file
)
from models import db, User, Draft
from config import DEFAULT_USER_PASSWORD, TEMPLATE_CONFIG
from routes.auth import admin_required, login_required
from services.document_service import (
    generate_documents_to_memory,
    get_cd_document_content,
    prepare_replacements
)
from services.pdf_service import (
    convert_all_docx_to_pdfs_batch,
    merge_pdfs_bytes,
    create_print_job,
    cleanup_print_job,
    cancel_print_job,
    is_job_cancelled
)

from helpers.html_helpers import (
    process_paragraph_html,
    process_table_html,
    generate_print_html_page
)
from docx import Document

logger = logging.getLogger(__name__)

admin_bp = Blueprint('admin', __name__)


@admin_bp.route('/admin/dashboard')
@admin_required
def admin_dashboard():
    """Admin dashboard"""
    user_id = session.get('user_id')
    admin = db.session.get(User, user_id)
    return render_template(
        'admin_dashboard.html',
        admin=admin,
        admin_name=admin.name
    )


@admin_bp.route('/api/admin/stats')
@admin_required
def api_admin_stats():
    """Get admin dashboard statistics"""
    try:
        total_users = User.query.filter_by(role='user').count()
        active_users = User.query.filter_by(role='user', is_active=True).count()
        total_drafts = Draft.query.count()
        draft_count = Draft.query.filter_by(status='draft').count()
        pending_count = Draft.query.filter_by(status='pending').count()
        approved_count = Draft.query.filter_by(status='approved').count()
        generated_count = Draft.query.filter_by(status='generated').count()

        users = User.query.filter_by(role='user').all()
        user_stats = []
        for user in users:
            user_stats.append({
                'id': user.id,
                'name': user.name,
                'email': user.email,
                'phone': user.phone,
                'is_active': user.is_active,
                'last_login': (
                    user.last_login.isoformat() if user.last_login else None
                ),
                'created_at': (
                    user.created_at.isoformat() if user.created_at else None
                ),
                'stats': {
                    'drafts': Draft.query.filter_by(
                        user_id=user.id, status='draft'
                    ).count(),
                    'pending': Draft.query.filter_by(
                        user_id=user.id, status='pending'
                    ).count(),
                    'approved': Draft.query.filter_by(
                        user_id=user.id, status='approved'
                    ).count(),
                    'generated': Draft.query.filter_by(
                        user_id=user.id, status='generated'
                    ).count(),
                    'total': Draft.query.filter_by(user_id=user.id).count()
                }
            })

        return jsonify({
            'success': True,
            'overall': {
                'total_users': total_users,
                'active_users': active_users,
                'total_documents': total_drafts,
                'drafts': draft_count,
                'pending': pending_count,
                'approved': approved_count,
                'generated': generated_count
            },
            'users': user_stats
        })

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/users', methods=['GET'])
@admin_required
def api_admin_get_users():
    """Get all users"""
    try:
        users = User.query.filter_by(role='user').order_by(
            User.created_at.desc()
        ).all()
        return jsonify({'success': True, 'users': [u.to_dict() for u in users]})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/users', methods=['POST'])
@admin_required
def api_admin_create_user():
    """Create a new user"""
    try:
        name = request.json.get('name', '').strip()
        email = request.json.get('email', '').strip().lower()
        phone = request.json.get('phone', '').strip()
        password = request.json.get('password', DEFAULT_USER_PASSWORD)
        role = request.json.get('role', 'user')

        if not name or not email:
            return jsonify({
                'success': False,
                'message': 'Name and email are required'
            })

        if not re.match(
            r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email
        ):
            return jsonify({'success': False, 'message': 'Invalid email format'})

        if User.get_by_email(email):
            return jsonify({'success': False, 'message': 'Email already exists'})

        new_user = User(
            name=name, email=email, phone=phone,
            role=role, is_active=True
        )
        new_user.set_password(password)
        db.session.add(new_user)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'User created successfully',
            'user': new_user.to_dict()
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/users/<user_id>', methods=['PUT'])
@admin_required
def api_admin_update_user(user_id):
    """Update a user"""
    try:
        user = db.session.get(User, user_id)
        if not user:
            return jsonify({'success': False, 'message': 'User not found'})

        name = request.json.get('name', '').strip()
        phone = request.json.get('phone', '').strip()
        new_email = request.json.get('email', '').strip().lower()
        new_password = request.json.get('password', '').strip()

        if name:
            user.name = name
        if phone is not None:
            user.phone = phone

        if new_email and new_email != user.email:
            existing = User.get_by_email(new_email)
            if existing and existing.id != user_id:
                return jsonify({'success': False, 'message': 'Email already exists'})
            user.email = new_email

        if new_password:
            user.set_password(new_password)

        db.session.commit()
        return jsonify({'success': True, 'message': 'User updated successfully'})

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/users/<user_id>/toggle', methods=['POST'])
@admin_required
def api_admin_toggle_user(user_id):
    """Toggle user active status"""
    try:
        user = db.session.get(User, user_id)
        if not user:
            return jsonify({'success': False, 'message': 'User not found'})

        user.is_active = not user.is_active
        db.session.commit()

        status = 'activated' if user.is_active else 'deactivated'
        return jsonify({
            'success': True,
            'message': f'User {status} successfully'
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/users/<user_id>', methods=['DELETE'])
@admin_required
def api_admin_delete_user(user_id):
    """Delete a user"""
    try:
        user = db.session.get(User, user_id)
        if not user:
            return jsonify({'success': False, 'message': 'User not found'})

        db.session.delete(user)
        db.session.commit()
        return jsonify({'success': True, 'message': 'User deleted successfully'})

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/users/<user_id>/approve', methods=['POST'])
@login_required
@admin_required
def api_approve_user(user_id):
    """Approve a user account"""
    try:
        current_user_id = session.get('user_id')
        current_user = db.session.get(User, current_user_id)

        if not current_user or not current_user.is_admin():
            return jsonify({'success': False, 'message': 'Unauthorized access'}), 403

        user = db.session.get(User, user_id)

        if not user:
            return jsonify({'success': False, 'message': 'User not found'}), 404

        if user.role in ['super_admin', 'admin']:
            return jsonify({
                'success': False,
                'message': 'Admin users are automatically approved'
            }), 400

        if user.is_approved:
            return jsonify({
                'success': False,
                'message': 'User is already approved'
            }), 400

        user.is_approved = True
        user.approved_at = datetime.now(timezone.utc)
        user.approved_by = current_user_id

        db.session.commit()

        logger.info(f'User approved: {user.email} by {current_user.email}')

        return jsonify({
            'success': True,
            'message': f'User "{user.name}" has been approved and can now login',
            'user': user.to_dict()
        }), 200

    except Exception as e:
        db.session.rollback()
        logger.error(f'Approve User Error: {str(e)}')
        return jsonify({'success': False, 'message': 'Failed to approve user'}), 500


@admin_bp.route('/api/admin/users/<user_id>/reject', methods=['POST'])
@login_required
@admin_required
def api_reject_user(user_id):
    """Reject/Revoke approval for a user account"""
    try:
        current_user_id = session.get('user_id')
        current_user = db.session.get(User, current_user_id)

        if not current_user or not current_user.is_admin():
            return jsonify({'success': False, 'message': 'Unauthorized access'}), 403

        user = db.session.get(User, user_id)

        if not user:
            return jsonify({'success': False, 'message': 'User not found'}), 404

        if user.role in ['super_admin', 'admin']:
            return jsonify({
                'success': False,
                'message': 'Admin users cannot be rejected'
            }), 400

        if not user.is_approved:
            return jsonify({
                'success': False,
                'message': 'User is already not approved'
            }), 400

        user.is_approved = False
        user.approved_at = None
        user.approved_by = None

        db.session.commit()

        logger.info(f'User approval revoked: {user.email} by {current_user.email}')

        return jsonify({
            'success': True,
            'message': f'Approval revoked for "{user.name}". They can no longer login.',
            'user': user.to_dict()
        }), 200

    except Exception as e:
        db.session.rollback()
        logger.error(f'Reject User Error: {str(e)}')
        return jsonify({'success': False, 'message': 'Failed to reject user'}), 500


@admin_bp.route('/api/admin/documents')
@admin_required
def api_admin_get_documents():
    """Get all documents for admin"""
    try:
        status = request.args.get('status')
        query = Draft.query.order_by(Draft.modified_at.desc())
        if status:
            query = query.filter_by(status=status)
        drafts = query.all()
        return jsonify({
            'success': True,
            'documents': [d.to_dict() for d in drafts]
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/documents/<doc_id>', methods=['DELETE'])
@login_required
def delete_document(doc_id):
    """Delete a document (admin only)"""
    try:
        if session.get('user_role') != 'admin':
            return jsonify({
                'success': False,
                'message': 'Admin access required'
            }), 403

        draft = db.session.get(Draft, doc_id)

        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'}), 404

        doc_name = draft.old_name or 'Untitled'

        try:
            if draft.status == 'generated' and draft.preview_data:
                output_folder = draft.preview_data.get('output_folder')
                if output_folder and Path(output_folder).exists():
                    import shutil
                    shutil.rmtree(output_folder, ignore_errors=True)
        except Exception as file_error:
            logger.warning(f'File cleanup error (non-critical): {file_error}')

        db.session.delete(draft)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Document "{doc_name}" deleted successfully'
        })

    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'Failed to delete document: {str(e)}'
        }), 500


@admin_bp.route('/api/admin/documents/<doc_id>', methods=['PUT'])
@admin_required
def api_admin_update_document(doc_id):
    """Admin update document"""
    try:
        draft = db.session.get(Draft, doc_id)
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'})

        data = request.json

        if 'replacements' in data:
            draft.replacements = data['replacements']
            draft.old_name = data['replacements'].get('OLD_NAME', draft.old_name)

        folder_type = data.get('folder_type', 'main')
        preview_data = draft.preview_data or {}
        template_type = draft.template_type

        if template_type in TEMPLATE_CONFIG:
            config = TEMPLATE_CONFIG[template_type]
            if folder_type == 'unmarried' and 'unmarried_subfolder' in config:
                preview_data['template_folder'] = config['unmarried_subfolder']
                preview_data['folder_type'] = 'unmarried'
            else:
                preview_data['template_folder'] = config['folder']
                preview_data['folder_type'] = 'main'

        draft.preview_data = preview_data
        draft.modified_at = datetime.now(timezone.utc)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Document updated successfully'})

    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/documents/<doc_id>/approve', methods=['POST'])
@admin_required
def api_admin_approve_document(doc_id):
    """Admin approve document"""
    try:
        draft = db.session.get(Draft, doc_id)
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'})

        draft.status = 'approved'
        draft.approved_at = datetime.now(timezone.utc)
        draft.modified_at = datetime.now(timezone.utc)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Document approved successfully'})

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/documents/<doc_id>/generate', methods=['POST'])
@admin_required
def api_admin_generate_document(doc_id):
    """Admin generate and download document"""
    try:
        draft = db.session.get(Draft, doc_id)
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'})

        if draft.status not in ['approved', 'generated']:
            return jsonify({
                'success': False,
                'message': 'Document must be approved first'
            })

        result = generate_documents_to_memory(draft, TEMPLATE_CONFIG)

        if not result['success']:
            return jsonify({
                'success': False,
                'message': result.get('message', 'Generation failed')
            })

        draft.status = 'generated'
        draft.generated_at = datetime.now(timezone.utc)
        draft.modified_at = datetime.now(timezone.utc)
        draft.generated_files = [f"Generated {result['file_count']} files"]
        db.session.commit()

        zip_buffer = result['zip_buffer']
        zip_buffer.seek(0)

        replacements_data = draft.replacements or {}
        if draft.template_type == 'minor_template' and replacements_data.get('FATHER-MOTHER_NAME'):
            safe_name = replacements_data['FATHER-MOTHER_NAME'].replace(' ', '_')
            filename = f"{safe_name}.zip"
        else:
            safe_name = draft.old_name.replace(' ', '_') if draft.old_name else 'document'
            filename = f"{safe_name}.zip"


        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/documents/<doc_id>/download', methods=['GET'])
@admin_required
def api_admin_download_document(doc_id):
    """Admin download already generated document"""
    try:
        draft = db.session.get(Draft, doc_id)
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'})

        result = generate_documents_to_memory(draft, TEMPLATE_CONFIG)
        if not result['success']:
            return jsonify({
                'success': False,
                'message': result.get('message', 'Generation failed')
            })

        zip_buffer = result['zip_buffer']
        zip_buffer.seek(0)

        replacements_data = draft.replacements or {}
        if draft.template_type == 'minor_template' and replacements_data.get('FATHER-MOTHER_NAME'):
            safe_name = replacements_data['FATHER-MOTHER_NAME'].replace(' ', '_')
            filename = f"{safe_name}.zip"
        else:
            safe_name = draft.old_name.replace(' ', '_') if draft.old_name else 'document'
            filename = f"{safe_name}.zip"


        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/documents/download-bulk', methods=['POST'])
@admin_required
def api_admin_download_bulk():
    """Admin download multiple documents as ZIP"""
    try:
        data = request.json
        doc_ids = data.get('doc_ids', [])

        if not doc_ids:
            return jsonify({'success': False, 'message': 'No documents selected'})

        bulk_zip_buffer = BytesIO()

        with zipfile.ZipFile(bulk_zip_buffer, 'w', zipfile.ZIP_DEFLATED) as bulk_zip:
            for doc_id in doc_ids:
                draft = db.session.get(Draft, doc_id)
                if not draft:
                    continue

                result = generate_documents_to_memory(draft, TEMPLATE_CONFIG)
                if not result['success']:
                    continue

                zip_buffer = result['zip_buffer']
                zip_buffer.seek(0)

                replacements_data = draft.replacements or {}
                if draft.template_type == 'minor_template' and replacements_data.get('FATHER-MOTHER_NAME'):
                    folder_name = replacements_data['FATHER-MOTHER_NAME'].replace(' ', '_')
                else:
                    folder_name = (
                        draft.old_name.replace(' ', '_')
                        if draft.old_name else f'doc_{doc_id}'
                    )


                with zipfile.ZipFile(zip_buffer, 'r') as individual_zip:
                    for file_info in individual_zip.filelist:
                        file_data = individual_zip.read(file_info.filename)
                        bulk_zip.writestr(
                            f'{folder_name}/{file_info.filename}', file_data
                        )

        bulk_zip_buffer.seek(0)

        return send_file(
            bulk_zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=(
                f"bulk_documents.zip"
            )
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})


@admin_bp.route('/api/admin/documents/<doc_id>/toggle-publish', methods=['POST'])
@admin_required
def api_toggle_publish(doc_id):
    """Toggle document published status"""
    try:
        draft = db.session.get(Draft, doc_id)
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'}), 404

        draft.published = not draft.published
        draft.published_at = datetime.now(timezone.utc) if draft.published else None
        draft.modified_at = datetime.now(timezone.utc)
        db.session.commit()

        status_text = 'Published' if draft.published else 'Unpublished'

        return jsonify({
            'success': True,
            'message': f'Document {status_text.lower()} successfully',
            'published': draft.published,
            'published_at': (
                draft.published_at.isoformat() if draft.published_at else None
            )
        })

    except Exception as e:
        db.session.rollback()
        logger.error(f'Toggle publish error: {str(e)}')
        return jsonify({'success': False, 'message': str(e)}), 500


@admin_bp.route('/api/admin/documents/<doc_id>/print-preview', methods=['GET'])
@admin_required
def api_print_preview(doc_id):
    """Get print preview for all documents"""
    try:
        draft = db.session.get(Draft, doc_id)

        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'})

        preview_data = draft.preview_data or {}
        template_folder = preview_data.get('template_folder')

        if not template_folder:
            template_config = TEMPLATE_CONFIG.get(draft.template_type, {})
            folder_type = preview_data.get('folder_type', 'main')

            if folder_type == 'unmarried' and 'unmarried_subfolder' in template_config:
                template_folder = template_config['unmarried_subfolder']
            else:
                template_folder = template_config.get('folder', '')

        if not template_folder:
            return jsonify({
                'success': False,
                'message': 'Template folder not configured'
            }), 400

        template_folder_path = Path(template_folder)
        replacements = prepare_replacements(draft.replacements or {})

        documents = []

        if template_folder_path.exists():
            for file in sorted(template_folder_path.iterdir()):
                if (file.is_file()
                        and file.suffix.lower() == '.docx'
                        and not file.name.startswith('~$')
                        and file.stem.upper() != 'CD'):    # ← exclude CD from progress log too
                    
                    
                    try:
                        doc = Document(str(file))
                        html_content = []

                        for para in doc.paragraphs:
                            text = para.text
                            for key, value in replacements.items():
                                if value is not None:
                                    text = text.replace(key, str(value))
                                else:
                                    text = text.replace(key, '')
                            import re
                            text = re.sub(r'\s+', ' ', text).strip()
                            if text:
                                html_content.append(f'<p>{text}</p>')

                        for table in doc.tables:
                            html_content.append(
                                '<table class="table table-bordered">'
                            )
                            for row in table.rows:
                                html_content.append('<tr>')
                                for cell in row.cells:
                                    cell_text = cell.text
                                    for key, value in replacements.items():
                                        if value is not None:
                                            cell_text = cell_text.replace(
                                                key, str(value)
                                            )
                                        else:
                                            cell_text = cell_text.replace(key, '')
                                    import re
                                    cell_text = re.sub(
                                        r'\s+', ' ', cell_text
                                    ).strip()
                                    html_content.append(f'<td>{cell_text}</td>')
                                html_content.append('</tr>')
                            html_content.append('</table>')

                        print_count = (
                            2 if file.stem.upper() == 'WITNESS' else 1
                        )

                        documents.append({
                            'filename': file.name,
                            'content': ''.join(html_content),
                            'print_count': print_count
                        })

                    except Exception as file_error:
                        logger.error(
                            f'Error processing {file.name}: {file_error}'
                        )
                        continue

        return jsonify({
            'success': True,
            'documents': documents,
            'document_name': draft.old_name or 'Unnamed'
        })

    except Exception as e:
        logger.error(f'Print preview error: {str(e)}')
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500


@admin_bp.route('/api/admin/documents/<doc_id>/print-html', methods=['GET'])
@admin_required
def api_print_html_preview(doc_id):
    """Generate HTML preview with formatting for printing"""
    try:
        draft = db.session.get(Draft, doc_id)

        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'}), 404

        preview_data = draft.preview_data or {}
        template_folder = preview_data.get('template_folder')

        if not template_folder:
            template_config = TEMPLATE_CONFIG.get(draft.template_type, {})
            folder_type = preview_data.get('folder_type', 'main')

            if folder_type == 'unmarried' and 'unmarried_subfolder' in template_config:
                template_folder = template_config['unmarried_subfolder']
            else:
                template_folder = template_config.get('folder', '')

        if not template_folder:
            return jsonify({
                'success': False,
                'message': 'Template folder not configured'
            }), 400

        template_folder_path = Path(template_folder)
        replacements = prepare_replacements(draft.replacements or {})

        clean_replacements = {
            key: html.escape(str(value)) if value is not None else ''
            for key, value in replacements.items()
        }

        all_documents_html = []

        docx_files = []
        if template_folder_path.exists():
            for file in sorted(template_folder_path.iterdir()):
                if (file.is_file()
                        and file.suffix.lower() == '.docx'
                        and not file.name.startswith('~$')):
                    docx_files.append(file)

        for docx_file in docx_files:
            doc = Document(str(docx_file))
            doc_html = []

            doc_html.append('''
            <div class="print-document-wrapper">
                <div class="print-document-body">
            ''')

            for paragraph in doc.paragraphs:
                para_html = process_paragraph_html(paragraph, clean_replacements)
                if para_html:
                    doc_html.append(para_html)

            for table in doc.tables:
                table_html = process_table_html(table, clean_replacements)
                if table_html:
                    doc_html.append(table_html)

            doc_html.append('</div></div>')
            all_documents_html.append(''.join(doc_html))

        full_html = generate_print_html_page(all_documents_html, draft.old_name)

        return full_html

    except Exception as e:
        logger.error(f'Print HTML error: {str(e)}')
        import traceback
        traceback.print_exc()
        return f"""
        <!DOCTYPE html>
        <html>
        <head><title>Error</title></head>
        <body>
            <div style="text-align: center; padding: 50px; color: red;">
                <h3>Error generating print preview</h3>
                <p>{html.escape(str(e))}</p>
            </div>
        </body>
        </html>
        """, 500


@admin_bp.route('/api/admin/documents/<doc_id>/print-pdf', methods=['GET'])
@admin_required
def api_print_pdf(doc_id):
    """Generate merged PDF — supports cancellation via job_id"""
    import time
    total_start = time.time()

    # ── Generate unique job_id and register it ───────────────────────
    job_id = uuid.uuid4().hex
    create_print_job(job_id)

    try:
        draft = db.session.get(Draft, doc_id)
        if not draft:
            cleanup_print_job(job_id)
            return jsonify({'success': False, 'message': 'Document not found'}), 404

        preview_data = draft.preview_data or {}
        template_folder = preview_data.get('template_folder')

        if not template_folder:
            template_config = TEMPLATE_CONFIG.get(draft.template_type, {})
            folder_type = preview_data.get('folder_type', 'main')
            if folder_type == 'unmarried' and 'unmarried_subfolder' in template_config:
                template_folder = template_config['unmarried_subfolder']
            else:
                template_folder = template_config.get('folder', '')

        if not template_folder:
            cleanup_print_job(job_id)
            return jsonify({
                'success': False,
                'message': 'Template folder not configured'
            }), 400

        template_folder_path = Path(template_folder)
        if not template_folder_path.exists():
            cleanup_print_job(job_id)
            return jsonify({
                'success': False,
                'message': f'Template folder not found: {template_folder}'
            }), 400

        replacements = prepare_replacements(draft.replacements or {})

        docx_files = sorted([
            f for f in template_folder_path.iterdir()
            if f.is_file()
            and f.suffix.lower() == '.docx'
            and not f.name.startswith('~$')
            and f.stem.upper() != 'CD'          # ← exclude CD.docx from printing

        ])

        if not docx_files:
            cleanup_print_job(job_id)
            return jsonify({
                'success': False,
                'message': 'No template files found'
            }), 400

        logger.info(
            f'[PrintPDF] Starting batch for doc_id={doc_id}, '
            f'job_id={job_id}, {len(docx_files)} templates'
        )

        docx_files_dict = {f.stem: f for f in docx_files}

        try:
            pdf_results = convert_all_docx_to_pdfs_batch(
                docx_files_dict, replacements, job_id=job_id
            )
        except RuntimeError as e:
            error_msg = str(e)
            cleanup_print_job(job_id)

            # ── Handle cancellation gracefully ───────────────────────
            if error_msg == 'CANCELLED':
                logger.info(f'[PrintPDF] Job {job_id} was cancelled by user')
                return jsonify({
                    'success': False,
                    'cancelled': True,
                    'message': 'Print job was cancelled'
                }), 200

            logger.error(f'[PrintPDF] Conversion failed: {error_msg}')

            if 'not found' in error_msg.lower():
                return jsonify({
                    'success': False,
                    'message': 'LibreOffice is not installed on this server.',
                    'error': error_msg
                }), 500
            else:
                return jsonify({
                    'success': False,
                    'message': f'PDF conversion failed: {error_msg}'
                }), 500

        pdf_list = []
        conversion_errors = []

        for docx_file in docx_files:
            stem = docx_file.stem
            if stem in pdf_results:
                pdf_bytes = pdf_results[stem]
                if stem.upper() == 'WITNESS':
                    pdf_list.append(pdf_bytes)
                    pdf_list.append(pdf_bytes)
                    logger.info(f'[PrintPDF] {stem}.pdf added ×2 (Witness)')
                else:
                    pdf_list.append(pdf_bytes)
                    logger.info(f'[PrintPDF] {stem}.pdf added ×1')
            else:
                conversion_errors.append(f'{stem}: PDF not generated')

        if not pdf_list:
            cleanup_print_job(job_id)
            return jsonify({
                'success': False,
                'message': 'No PDFs were generated successfully',
                'errors': conversion_errors
            }), 500

        if len(pdf_list) == 1:
            merged_pdf = pdf_list[0]
        else:
            try:
                merge_start = time.time()
                merged_pdf = merge_pdfs_bytes(pdf_list)
                logger.info(
                    f'[PrintPDF] Merged {len(pdf_list)} PDFs '
                    f'in {time.time() - merge_start:.2f}s'
                )
            except ImportError:
                logger.warning('[PrintPDF] pypdf not installed')
                merged_pdf = pdf_list[0]
            except Exception as e:
                logger.error(f'[PrintPDF] Merge error: {e}')
                merged_pdf = pdf_list[0]

        total_elapsed = time.time() - total_start
        logger.info(
            f'[PrintPDF] ✅ Total time: {total_elapsed:.2f}s '
            f'for {len(docx_files)} files'
        )

        safe_name = (draft.old_name or 'document').replace(' ', '_')
        filename = (
            f"Print_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        )

        response = send_file(
            BytesIO(merged_pdf),
            mimetype='application/pdf',
            as_attachment=False,
            download_name=filename
        )

        # ── Expose job_id in header so frontend can cancel ───────────
        response.headers['X-Print-Job-Id'] = job_id
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        logger.error(f'[PrintPDF] Fatal error: {str(e)}')
        cleanup_print_job(job_id)
        return jsonify({'success': False, 'message': str(e)}), 500

    finally:
        cleanup_print_job(job_id)

@admin_bp.route('/api/admin/documents/<doc_id>/cancel-print', methods=['POST'])
@admin_required
def api_cancel_print(doc_id):
    """Cancel an in-progress print job"""
    try:
        data = request.get_json() or {}
        job_id = data.get('job_id')

        if not job_id:
            return jsonify({
                'success': False,
                'message': 'No job_id provided'
            }), 400

        success, message = cancel_print_job(job_id)

        if not success:
            return jsonify({'success': False, 'message': message}), 404

        logger.info(f'[CancelPrint] Job {job_id} cancelled for doc_id={doc_id}')
        return jsonify({'success': True, 'message': message})

    except Exception as e:
        logger.error(f'[CancelPrint] Error: {str(e)}')
        return jsonify({'success': False, 'message': str(e)}), 500

@admin_bp.route('/api/admin/documents/<doc_id>/download-all', methods=['GET'])
@admin_required
def api_download_all_files(doc_id):
    """Download all document files as ZIP"""
    try:
        draft = db.session.get(Draft, doc_id)

        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'}), 404

        result = generate_documents_to_memory(draft, TEMPLATE_CONFIG)

        if not result['success']:
            return jsonify({
                'success': False,
                'message': result.get('message', 'Generation failed')
            }), 500

        zip_buffer = result['zip_buffer']
        zip_buffer.seek(0)

        replacements_data = draft.replacements or {}
        if draft.template_type == 'minor_template' and replacements_data.get('FATHER-MOTHER_NAME'):
            safe_name = replacements_data['FATHER-MOTHER_NAME'].replace(' ', '_')
            filename = f"{safe_name}.zip"
        else:
            safe_name = draft.old_name.replace(' ', '_') if draft.old_name else 'document'
            filename = f"{safe_name}.zip"

        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        logger.error(f'Download all error: {str(e)}')
        return jsonify({'success': False, 'message': str(e)}), 500


@admin_bp.route('/api/admin/documents/<doc_id>/download-cd', methods=['GET'])
@admin_required
def api_download_cd_only(doc_id):
    """Download CD document only"""
    try:
        draft = db.session.get(Draft, doc_id)
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'}), 404

        preview_data = draft.preview_data or {}
        template_folder = preview_data.get('template_folder')

        if not template_folder:
            template_config = TEMPLATE_CONFIG.get(draft.template_type, {})
            folder_type = preview_data.get('folder_type', 'main')
            if folder_type == 'unmarried' and 'unmarried_subfolder' in template_config:
                template_folder = template_config['unmarried_subfolder']
            else:
                template_folder = template_config.get('folder', '')

        if not template_folder:
            return jsonify({
                'success': False,
                'message': 'Template folder not configured'
            }), 400

        template_folder_path = Path(template_folder)

        cd_file = None
        if template_folder_path.exists():
            for file in template_folder_path.iterdir():
                if (file.is_file()
                        and file.suffix.lower() == '.docx'
                        and file.stem.upper() == 'CD'
                        and not file.name.startswith('~$')):
                    cd_file = file
                    break

        if not cd_file:
            return jsonify({'success': False, 'message': 'CD.docx not found'}), 404

        doc = Document(str(cd_file))
        replacements = prepare_replacements(draft.replacements or {})

        from helpers.docx_helpers import (
            replace_text_in_paragraph,
            replace_text_in_tables
        )

        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)

        if doc.tables:
            replace_text_in_tables(doc.tables, replacements)

        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            if section.header.tables:
                replace_text_in_tables(section.header.tables, replacements)
            for paragraph in section.footer.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            if section.footer.tables:
                replace_text_in_tables(section.footer.tables, replacements)

        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        replacements_data = draft.replacements or {}
        if draft.template_type == 'minor_template' and replacements_data.get('FATHER-MOTHER_NAME'):
            safe_name = replacements_data['FATHER-MOTHER_NAME'].replace(' ', '_')
            filename = f"CD_{safe_name}.docx"
        else:
            safe_name = draft.old_name.replace(' ', '_') if draft.old_name else 'CD'
            filename = f"CD_{safe_name}.docx"
        

        return send_file(
            doc_buffer,
            mimetype=(
                'application/vnd.openxmlformats-officedocument'
                '.wordprocessingml.document'
            ),
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        logger.error(f'Download CD error: {str(e)}')
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500


@admin_bp.route('/api/admin/documents/<doc_id>/cd-preview', methods=['GET'])
@admin_required
def api_get_cd_preview(doc_id):
    """Get CD preview for admin"""
    try:
        draft = db.session.get(Draft, doc_id)
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'}), 404

        preview_data = draft.preview_data or {}
        template_folder = preview_data.get('template_folder')

        if not template_folder:
            template_config = TEMPLATE_CONFIG.get(draft.template_type, {})
            folder_type = preview_data.get('folder_type', 'main')
            if folder_type == 'unmarried' and 'unmarried_subfolder' in template_config:
                template_folder = template_config['unmarried_subfolder']
            else:
                template_folder = template_config.get('folder', '')

        if not template_folder:
            return jsonify({
                'success': False,
                'message': 'Template folder not configured'
            }), 400

        replacements = prepare_replacements(draft.replacements or {})
        cd_content = get_cd_document_content(template_folder, replacements)

        return jsonify({
            'success': True,
            'cd_content': cd_content,
            'document_name': draft.old_name or 'Unnamed'
        })

    except Exception as e:
        logger.error(f'Error getting CD preview: {str(e)}')
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500